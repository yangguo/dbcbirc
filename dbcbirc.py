import datetime
import glob
import io
import json
import os
import random
import re
import time
from ast import literal_eval

import docx
import numpy as np
import pandas as pd
import requests
import streamlit as st
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt
from pyecharts.render import make_snapshot
import plotly.express as px

import snapshot as driver
from database import delete_data, get_collection, get_data, insert_data
from utils import split_words
from collections import Counter
# from streamlit_tags import st_tags


# import matplotlib

mappath = "map/china.json"
pencbirc = "cbirc"
# mapfolder = '../temp/citygeo.csv'
urldtl = "https://www.cbirc.gov.cn/cn/view/pages/ItemDetail.html?docId="
# choose orgname index
org2name = {
    "银保监会机关": "jiguan",
    "银保监局本级": "benji",
    "银保监分局本级": "fenju",
    "": "",
}

city2province = {
    "浙江省": "浙江省",
    "河南省": "河南省",
    "安徽省": "安徽省",
    "青海省": "青海省",
    "陕西省": "陕西省",
    "深圳": "广东省",
    "黑龙江省": "黑龙江省",
    "湖南省": "湖南省",
    "山东省": "山东省",
    "云南省": "云南省",
    "甘肃省": "甘肃省",
    "吉林省": "吉林省",
    "北京": "北京市",
    "内蒙古自治区": "内蒙古自治区",
    "江西省": "江西省",
    "新疆维吾尔自治区": "新疆维吾尔自治区",
    "福建省": "福建省",
    "江苏省": "江苏省",
    "四川省": "四川省",
    "贵州省": "贵州省",
    "广西壮族自治区": "广西壮族自治区",
    "广东省": "广东省",
    "山西省": "山西省",
    "海南省": "海南省",
    "湖北省": "湖北省",
    "无": "未知省份",
    "全国": "未知省份",
    "大连": "辽宁省",
    "辽宁省": "辽宁省",
    "新疆": "新疆维吾尔自治区",
    "重庆市": "重庆市",
    "上海市": "上海市",
    "天津市": "天津市",
    "宁夏回族自治区": "宁夏回族自治区",
    "西藏自治区": "西藏自治区",
    "河北省": "河北省",
    "深圳市": "广东省",
    "北京市": "北京市",
    "贵州省黔南布依族苗族自治州": "贵州省",
    "天津": "天津市",
    "新疆哈密": "新疆维吾尔自治区",
    "新疆哈密市": "新疆维吾尔自治区",
    "新疆维吾尔自治区哈密市": "新疆维吾尔自治区",
    "新疆维吾尔自治区昌吉回族自治州": "新疆维吾尔自治区",
    "新疆维吾尔自治区博尔塔拉蒙古自治州": "新疆维吾尔自治区",
    "甘肃省甘南藏族自治州": "甘肃省",
    "山东省威海市": "山东省",
    "四川省绵阳市": "四川省",
    "内蒙古自治区包头市": "内蒙古自治区",
    "新疆石河子": "新疆维吾尔自治区",
    "内蒙古自治区乌海市": "内蒙古自治区",
    "新疆石河子市": "新疆维吾尔自治区",
    "云南省楚雄彝族自治州": "云南省",
    "广东省河源市": "广东省",
    "上海": "上海市",
    "青岛市": "山东省",
    "湖南省湘潭市": "湖南省",
    "贵州省遵义市": "贵州省",
    "山东省菏泽市": "山东省",
    "江苏省淮安市": "江苏省",
    "商洛市": "陕西省",
    "重庆": "重庆市",
    "汉中": "陕西省",
    "江西省萍乡市": "江西省",
    "江苏省盐城市": "江苏省",
    "辽宁省锦州市": "辽宁省",
    "湖南省益阳市": "湖南省",
    "新疆维吾尔自治区和田地区": "新疆维吾尔自治区",
    "葫芦岛市": "辽宁省",
    "河南省开封市": "河南省",
    "辽宁省葫芦岛市": "辽宁省",
    "广东省揭阳市": "广东省",
    "新疆维吾尔自治区吐鲁番市": "新疆维吾尔自治区",
    "本溪市": "辽宁省",
    "广东省佛山市": "广东省",
    "晋城市": "山西省",
    "新疆和田": "新疆维吾尔自治区",
    "安徽省淮南市": "安徽省",
    "安徽省宣城市": "安徽省",
    "宣城市": "安徽省",
    "江西省抚州市": "江西省",
    "黑龙江省大庆市": "黑龙江省",
    "广西壮族自治区来宾市": "广西壮族自治区",
    "四川省达州市": "四川省",
    "喀什": "新疆维吾尔自治区",
    "海西": "青海省",
    "浙江省衢州市": "浙江省",
    "湖南省邵阳市": "湖南省",
    "新疆维吾尔自治区阿勒泰地区": "新疆维吾尔自治区",
    "广西壮族自治区钦州市": "广西壮族自治区",
    "黔南州": "贵州省",
    "大兴安岭地区": "黑龙江省",
    "江苏省苏州市": "江苏省",
    "湖南省岳阳市": "湖南省",
    "河南省许昌市": "河南省",
    "许昌市": "河南省",
    "安徽省阜阳市": "安徽省",
    "湖南省郴州市": "湖南省",
    "湖南省株洲市": "湖南省",
    "中山市": "广东省",
    "张家界市": "湖南省",
    "吕梁市": "山西省",
    "四川省南充市": "四川省",
    "河南省信阳市": "河南省",
    "乌兰察布市": "内蒙古自治区",
    "广东省中山市": "广东省",
    "新疆维吾尔自治区克孜勒苏柯尔克孜自治州": "新疆维吾尔自治区",
    "吐鲁番市": "新疆维吾尔自治区",
    "新疆博尔塔拉蒙古自治州": "新疆维吾尔自治区",
    "新疆阿勒泰地区": "新疆维吾尔自治区",
    "大庆市": "黑龙江省",
    "福建省泉州市": "福建省",
    "新疆巴音郭楞蒙古自治州": "新疆维吾尔自治区",
    "广西壮族自治区贵港市": "广西壮族自治区",
    "海西地区": "山东省",
    "广西贵港市": "广西壮族自治区",
    "内蒙古自治区阿拉善盟": "内蒙古自治区",
    "内蒙古自治区乌兰察布市": "内蒙古自治区",
    "攀枝花市": "四川省",
    "承德市": "河北省",
    "广东省云浮市": "广东省",
    "鞍山": "辽宁省",
    "绍兴市": "浙江省",
    "湖北省鄂州市": "湖北省",
    "威海市": "山东省",
    "克孜勒苏": "新疆维吾尔自治区",
    "江苏省常州市": "江苏省",
    "河南省周口市": "河南省",
    "四川省甘孜州": "四川省",
    "安徽省蚌埠市": "安徽省",
    "贵州省六盘水市": "贵州省",
    "南充市": "四川省",
    "云南省德宏州": "云南省",
    "陕西省渭南市": "陕西省",
    "宁夏回族自治区吴忠市": "宁夏回族自治区",
    "浙江省舟山市": "浙江省",
    "苏州": "江苏省",
    "广西壮族自治区柳州市": "广西壮族自治区",
    "海南省三亚市": "海南省",
    "张掖市": "甘肃省",
    "新疆维吾尔自治区塔城地区": "新疆维吾尔自治区",
    "塔城地区": "新疆维吾尔自治区",
    "云南省曲靖市": "云南省",
    "云南省普洱市": "云南省",
    "新疆维吾尔自治区阿克苏地区": "新疆维吾尔自治区",
    "湖南省衡阳市": "湖南省",
    "广东省江门市": "广东省",
    "云南省丽江市": "云南省",
    "济宁市": "山东省",
    "湖南省怀化市": "湖南省",
    "浙江省金华市": "浙江省",
    "金华市": "浙江省",
    "山东省烟台市": "山东省",
    "吐鲁番": "新疆维吾尔自治区",
    "广东省潮州市": "广东省",
    "浙江省绍兴市": "浙江省",
    "淄博市": "山东省",
    "广西来宾": "广西壮族自治区",
    "驻马店市": "河南省",
    "聊城市": "山东省",
    "滁州市": "安徽省",
    "果洛": "青海省",
    "海北州": "青海省",
    "浙江省丽水市": "浙江省",
    "丽水市": "浙江省",
    "贵港市": "广西壮族自治区",
    "玉树": "青海省",
    "滨州市": "山东省",
    "大兴安岭": "内蒙古自治区",
    "黑龙江省伊春市": "黑龙江省",
    "鸡西市": "黑龙江省",
    "辽宁省本溪市": "辽宁省",
    "鄂州市": "湖北省",
    "云南省文山州": "云南省",
    "广东省梅州市": "广东省",
    "菏泽市": "山东省",
    "吐鲁番地区": "新疆维吾尔自治区",
    "乌海": "内蒙古自治区",
    "新疆克孜勒苏柯尔克孜自治州": "新疆维吾尔自治区",
    "陕西省汉中市": "陕西省",
    "新疆塔城": "新疆维吾尔自治区",
    "辽宁省鞍山市": "辽宁省",
    "贵州省黔南州": "贵州省",
    "抚顺市": "辽宁省",
    "辽源市": "吉林省",
    "江西省鹰潭市": "江西省",
    "江苏省连云港市": "江苏省",
    "浙江省宁波市": "浙江省",
    "厦门": "福建省",
    "宁波市": "浙江省",
    "宁波": "浙江省",
    "青岛": "山东省",
    "广西": "广西壮族自治区",
    "辽宁省大连市": "辽宁省",
    "福建省厦门市": "福建省",
    "大连市": "辽宁省",
    "厦门市": "福建省",
    "内蒙古鄂尔多斯市": "内蒙古自治区",
    "石家庄市": "河北省",
    "中国": "甘肃省",
    "江西省上饶市": "江西省",
    "四川省眉山市": "四川省",
    "广东省东莞市": "广东省",
    "河北省邢台市": "河北省",
    "江苏省镇江市": "江苏省",
    "广东省肇庆市": "广东省",
}


# Normalize various region strings to standard province names used by the map
def normalize_province_name(name: str) -> str:
    # Guard against non-string or empty inputs
    try:
        region = str(name).strip()
    except Exception:
        return "未知省份"

    if not region:
        return "未知省份"

    # First try explicit exact mapping table
    if region in city2province:
        return city2province[region]

    # Handle common province/municipality/autonomous region keywords
    shorthand_map = {
        "北京": "北京市",
        "天津": "天津市",
        "上海": "上海市",
        "重庆": "重庆市",
        "河北": "河北省",
        "山西": "山西省",
        "辽宁": "辽宁省",
        "吉林": "吉林省",
        "黑龙江": "黑龙江省",
        "江苏": "江苏省",
        "浙江": "浙江省",
        "安徽": "安徽省",
        "福建": "福建省",
        "江西": "江西省",
        "山东": "山东省",
        "河南": "河南省",
        "湖北": "湖北省",
        "湖南": "湖南省",
        "广东": "广东省",
        "海南": "海南省",
        "四川": "四川省",
        "贵州": "贵州省",
        "云南": "云南省",
        "陕西": "陕西省",
        "甘肃": "甘肃省",
        "青海": "青海省",
        "内蒙古": "内蒙古自治区",
        "广西": "广西壮族自治区",
        "西藏": "西藏自治区",
        "宁夏": "宁夏回族自治区",
        "新疆": "新疆维吾尔自治区",
    }
    # If region CONTAINS any of the province keywords, map accordingly
    for keyword, std_province in shorthand_map.items():
        if keyword in region:
            return std_province

    # If looks like a full province name already, return as-is
    if region.endswith("省") or region.endswith("市") or region.endswith("自治区"):
        return region

    # Try matching by known city/area keywords from the mapping table
    # Prefer longer (more specific) keys first to avoid partial mismatches
    for key in sorted(city2province.keys(), key=len, reverse=True):
        # Avoid overly generic keys to prevent false positives
        if key in ("中国",):
            continue
        if key and key in region:
            return city2province[key]

    # Fallback
    return "未知省份"

# @st.cache(allow_output_mutation=True)
def get_csvdf(penfolder, beginwith):
    # read parquet file if exists
    filepath = penfolder + "/" + beginwith + ".parquet"
    # if os.path.exists(filepath):
    #     df = pd.read_parquet(filepath)
    # else:
    
    # 只读取根目录下的文件，不递归查找子目录
    files2 = glob.glob(os.path.join(penfolder, beginwith + "*.csv"))
    
    dflist = []
    # filelist = []
    for filepath in files2:
        try:
            pendf = pd.read_csv(filepath)
            dflist.append(pendf)
            # filelist.append(filename)
        except Exception as e:
            print(f"读取文件 {filepath} 失败: {e}")
    
    if len(dflist) > 0:
        df = pd.concat(dflist)
        df.reset_index(drop=True, inplace=True)
    else:
        df = pd.DataFrame()
        # convert df to paquet
        # df.to_parquet(penfolder + "/" + beginwith + ".parquet")
    return df


# read parquet
def get_parquetdf(penfolder, beginwith):
    filepath = penfolder + "/" + beginwith + ".parquet"
    df = pd.read_parquet(filepath)
    return df


# @st.cache(allow_output_mutation=True)
def get_cbircanalysis(orgname):
    org_name_index = org2name[orgname]
    beginwith = "cbircsplit" + org_name_index
    # beginwith = "tempsplit"
    pendf = get_csvdf(pencbirc, beginwith)
    # if pendf is not empty
    # if len(pendf) > 0:
    #     # format date
    #     # pendf["发布日期"] = pd.to_datetime(pendf["发布日期"]).dt.date
    #     # fillna
    #     pendf.fillna("", inplace=True)
    # else:
    #     pendf = pd.DataFrame()
    return pendf


# @st.cache(allow_output_mutation=True)
def get_cbircdetail(orgname):
    org_name_index = org2name[orgname]

    beginwith = "cbircdtl" + org_name_index
    d0 = get_csvdf(pencbirc, beginwith)
    # Check if DataFrame is empty
    if d0.empty:
        st.warning(f"No detail data available for {orgname}")
        return pd.DataFrame()
    
    # Check if required columns exist
    required_columns = ["title", "subtitle", "date", "doc", "id"]
    missing_columns = [col for col in required_columns if col not in d0.columns]
    if missing_columns:
        st.warning(f"Missing columns in detail data for {orgname}: {missing_columns}")
        return pd.DataFrame()
    
    # reset index
    d1 = d0[["title", "subtitle", "date", "doc", "id"]].reset_index(drop=True)
    # format date
    # d1["date"] = pd.to_datetime(d1["date"]).dt.date
    d1["date"] = d1["date"].str.split(".").str[0]
    d1["date"] = pd.to_datetime(d1["date"], format="%Y-%m-%d %H:%M:%S").dt.date
    # update column name
    d1.columns = ["标题", "文号", "发布日期", "内容", "id"]
    # fillna
    # d1.fillna("", inplace=True)
    return d1


def display_cbircsum(df):
    # get length of old eventdf
    oldlen = len(df)
    # Check if DataFrame is empty or if required column exists
    if df.empty or "发布日期" not in df.columns:
        col1, col2 = st.columns([1, 3])
        col1.metric("案例总数", oldlen)
        col2.metric("日期范围", "无数据")
        return
    
    # get min and max date of old eventdf
    min_date = df["发布日期"].min()
    max_date = df["发布日期"].max()
    # use metric for length and date
    col1, col2 = st.columns([1, 3])
    col1.metric("案例总数", oldlen)
    col2.metric("日期范围", f"{min_date} - {max_date}")


def get_cbircsum(orgname):
    org_name_index = org2name[orgname]
    beginwith = "cbircsum" + org_name_index
    pendf = get_csvdf(pencbirc, beginwith)
    # Check if DataFrame is empty or if required column exists
    if pendf.empty:
        st.warning(f"No summary data available for {orgname}")
        return pd.DataFrame()
    if "publishDate" not in pendf.columns:
        st.warning(f"Column 'publishDate' not found in summary data for {orgname}")
        return pendf
    # format date
    pendf["发布日期"] = pd.to_datetime(pendf["publishDate"]).dt.date
    return pendf


def get_cbirctoupd(orgname):
    org_name_index = org2name[orgname]
    beginwith = "cbirctoupd" + org_name_index
    pendf = get_csvdf(pencbirc, beginwith)
    return pendf


def searchcbirc(
    df,
    start_date,
    end_date,
    wenhao_text,
    people_text,
    event_text,
    law_text,
    penalty_text,
    org_text,
    industry,
    min_penalty,
    province,
):
    # cols = [
    #     "标题",
    #     "文号",
    #     "发布日期",
    #     "行政处罚决定书文号",
    #     "被处罚当事人",
    #     "主要违法违规事实",
    #     "行政处罚依据",
    #     "行政处罚决定",
    #     "作出处罚决定的机关名称",
    #     "作出处罚决定的日期",
    #     "id",
    #     "label",
    #     "amount",
    #     "province",
    # ]
    # split words
    wenhao_text = split_words(wenhao_text)
    people_text = split_words(people_text)
    event_text = split_words(event_text)
    # law_text = split_words(law_text)
    penalty_text = split_words(penalty_text)
    org_text = split_words(org_text)

    # search by start_date and end_date, wenhao_text, people_text, event_text, law_text, penalty_text, org_text
    searchdf = df[
        (df["发布日期"] >= start_date)
        & (df["发布日期"] <= end_date)
        # & (df["行政处罚决定书文号"].str.contains(wenhao_text))
        & (df["wenhao"].str.contains(wenhao_text))
        # & (df["被处罚当事人"].str.contains(people_text))
        & (df["people"].str.contains(people_text))
        # & (df["主要违法违规事实"].str.contains(event_text))
        & (df["event"].str.contains(event_text))
        # & (df["行政处罚依据"].str.contains(law_text))
        # & (df["行政处罚决定"].str.contains(penalty_text))
        & (df["penalty"].str.contains(penalty_text))
        # & (df["作出处罚决定的机关名称"].str.contains(org_text))
        & (df["org"].str.contains(org_text))
        # & (df["label"].isin(industry))
        & (df["industry"].str.contains(industry))
        & (df["amount"] >= min_penalty)
        # & (df["法律法规"].isin(law_text))
        & (df["law"].str.contains(law_text))
        # & (df["province"].isin(province))
        & (df["province"].str.contains(province))
    ]  # [cols]
    # sort by date desc
    # searchdf.sort_values(by=["发布日期"], ascending=False, inplace=True)
    # drop duplicates
    # searchdf.drop_duplicates(subset=["id"], inplace=True)
    searchdf = searchdf.loc[:].sort_values(by=["发布日期"], ascending=False)
    searchdf = searchdf.loc[:].drop_duplicates(subset=["id"])
    # reset index
    searchdf = searchdf.reset_index(drop=True)
    return searchdf


# search by title, subtitle, date, doc, industry
def searchdtl(
    df,
    start_date,
    end_date,
    title_text,
    wenhao_text,
    event_text,
    industry,
    min_penalty,
    law_select,
    province_select,
):
    # cols = ["标题", "文号", "发布日期", "内容", "id", "label", "amount", "province"]
    # split words
    title_text = split_words(title_text)
    wenhao_text = split_words(wenhao_text)
    event_text = split_words(event_text)
    # search by start_date and end_date, wenhao_text, people_text, event_text, law_text, penalty_text, org_text
    # st.write(df[:5000])
    # set index on label, province, 发布日期, and amount
    # indexed_df = df.set_index(["label", "province", "发布日期", "amount"])
    # # search by start_date and end_date, wenhao_text, people_text, event_text, law_text, penalty_text, org_text
    # searchdf = indexed_df.loc[
    #     (
    #         slice(None),
    #         slice(None),
    #         slice(start_date, end_date),
    #         slice(min_penalty, None),
    #     ),
    #     cols,
    # ]
    # searchdf = searchdf[
    #     searchdf["标题"].str.contains(title_text)
    #     & searchdf["文号"].str.contains(wenhao_text)
    #     & searchdf["内容"].str.contains(event_text)
    #     & searchdf["label"].isin(industry)
    #     & searchdf["法律法规"].isin(law_select)
    #     & searchdf["province"].isin(province_select)
    # ]
    searchdf = df[
        (df["发布日期"] >= start_date)
        & (df["发布日期"] <= end_date)
        & (df["标题"].str.contains(title_text))
        & (df["文号"].str.contains(wenhao_text))
        & (df["内容"].str.contains(event_text))
        # & (df["label"].isin(industry))
        & (df["industry"].str.contains(industry))
        & (df["amount"] >= min_penalty)
        # & (df["法律法规"].isin(law_select))
        & (df["law"].str.contains(law_select))
        # & (df["province"].isin(province_select))
        & (df["province"].str.contains(province_select))
    ]  # [cols]
    # sort by date desc
    searchdf.sort_values(by=["发布日期"], ascending=False, inplace=True)
    # drop duplicates by id
    searchdf.drop_duplicates(subset=["id"], inplace=True)
    # reset index
    searchdf.reset_index(drop=True, inplace=True)
    return searchdf


# count the number of df by month
def count_by_month(df):
    df_month = df.copy()
    # count by month
    df_month["month"] = df_month["发布日期"].apply(lambda x: x.strftime("%Y-%m"))
    df_month_count = df_month.groupby(["month"]).size().reset_index(name="count")
    return df_month_count


# display dfmonth in plotly
def display_dfmonth(df):
    df_month = df.copy()
    # count by month
    df_month["month"] = df_month["发布日期"].apply(lambda x: x.strftime("%Y-%m"))
    df_month_count = df_month.groupby(["month"]).size().reset_index(name="count")
    # display checkbox to show/hide graph1
    # showgraph1 = st.sidebar.checkbox("按发文时间统计", key="showgraph1")
    showgraph1 = True
    if showgraph1:
        x_data = df_month_count["month"].tolist()
        y_data = df_month_count["count"].tolist()
        # draw echarts bar chart
        # bar = (
        #     Bar()
        #     .add_xaxis(xaxis_data=x_data)
        #     .add_yaxis(series_name="数量", y_axis=y_data, yaxis_index=0)
        #     .set_global_opts(
        #         title_opts=opts.TitleOpts(title="按发文时间统计"),
        #         visualmap_opts=opts.VisualMapOpts(max_=max(y_data), min_=min(y_data)),
        #     )
        # )
        # # use events
        # events = {
        #     "click": "function(params) { console.log(params.name); return params.name }",
        #     # "dblclick":"function(params) { return [params.type, params.name, params.value] }"
        # }
        # use events
        # yearmonth = st_pyecharts(bar, events=events)
        bar, yearmonth = print_bar(x_data, y_data, "处罚数量", "按发文时间统计")
        # st.write(yearmonth)
        if yearmonth is not None:
            # get year and month value from format "%Y-%m"
            # year = int(yearmonth.split("-")[0])
            # month = int(yearmonth.split("-")[1])
            # filter date by year and month
            searchdfnew = df_month[df_month["month"] == yearmonth]
            # drop column "month"
            searchdfnew.drop(columns=["month"], inplace=True)

            # set session state
            st.session_state["search_result_cbirc"] = searchdfnew

        # 图一解析开始
        maxmonth = df_month["month"].max()
        minmonth = df_month["month"].min()
        # get total number of count
        num_total = len(df_month["month"])
        # get total number of month count
        month_total = len(set(df_month["month"].tolist()))
        # get average number of count per month count
        num_avg = num_total / month_total
        # get month value of max count
        top1month = max(
            set(df_month["month"].tolist()), key=df_month["month"].tolist().count
        )
        top1number = df_month["month"].tolist().count(top1month)

        image1_text = (
            "图一解析：从"
            + minmonth
            + "至"
            + maxmonth
            + "，共发生"
            + str(num_total)
            + "起处罚事件，"
            + "平均每月发生"
            + str(round(num_avg))
            + "起处罚事件。其中"
            + top1month
            + "最高发生"
            + str(top1number)
            + "起处罚事件。"
        )

        # display total coun
        st.markdown("##### " + image1_text)

    # get eventdf sum amount by month
    df_sum, df_sigle_penalty = sum_amount_by_month(df_month)

    sum_data = df_sum["sum"].tolist()
    line, yearmonthline = print_line(x_data, sum_data, "处罚金额", "案例金额统计")

    if yearmonthline is not None:
        # filter date by year and month
        searchdfnew = df_month[df_month["month"] == yearmonthline]
        # drop column "month"
        searchdfnew.drop(columns=["month"], inplace=True)
        # set session state
        st.session_state["search_result_cbirc"] = searchdfnew
        # refresh page
        # st.experimental_rerun()

    # 图二解析：
    sum_data_number = 0  # 把案件金额的数组进行求和
    more_than_100 = 0  # 把案件金额大于100的数量进行统计
    case_total = 0  # 把案件的总数量进行统计
    penaltycount = df_sigle_penalty["amount"].tolist()
    for i in penaltycount:
        sum_data_number = sum_data_number + i / 10000
        if i > 100 * 10000:
            more_than_100 = more_than_100 + 1
        if i != 0:
            case_total = case_total + 1

    # for i in sum_data:
    #     sum_data_number = sum_data_number + i / 10000
    #     if i > 100 * 10000:
    #         more_than_100 = more_than_100 + 1
    # sum_data_number=round(sum_data_number,2)
    if case_total > 0:
        avg_sum = round(sum_data_number / case_total, 2)
    else:
        avg_sum = 0
    # get index of max sum
    topsum1 = df_sum["sum"].nlargest(1)
    topsum1_index = df_sum["sum"].idxmax()
    # get month value of max count
    topsum1month = df_sum.loc[topsum1_index, "month"]
    image2_text = (
        "图二解析：从"
        + minmonth
        + "至"
        + maxmonth
        + "，共发生罚款案件"
        + str(case_total)
        + "起;期间共涉及处罚金额"
        + str(round(sum_data_number, 2))
        + "万元，处罚事件平均处罚金额为"
        + str(avg_sum)
        + "万元，其中处罚金额高于100万元处罚事件共"
        + str(more_than_100)
        + "起。"
        + topsum1month
        + "发生最高处罚金额"
        + str(round(topsum1.values[0] / 10000, 2))
        + "万元。"
    )
    st.markdown("##### " + image2_text)

    # locdf = get_cbircloc()
    # merge df_month with locdf by id
    # df_month_loc = pd.merge(df_month, locdf, on="id", how="left")
    df_month_loc = df_month
    # count by province
    df_org_count = df_month_loc.groupby(["province"]).size().reset_index(name="count")
    # sort by count
    df_org_count = df_org_count.sort_values(by="count", ascending=False)
    org_ls = df_org_count["province"].tolist()
    count_ls = df_org_count["count"].tolist()

    new_orgls, new_countls = count_by_province(org_ls, count_ls)
    map_data = print_map(new_orgls, new_countls, "处罚地图")
    # st_pyecharts(map_data, map=map, width=800, height=650)
    # display map
    # components.html(map_data.render_embed(), height=650)
    # image3_text = "图三解析：处罚地图"

    pie, orgname = print_pie(new_orgls, new_countls, "按发文机构统计")

    # 图四解析开始
    # orgls = pd.value_counts(df_month_loc["province"]).keys().tolist()
    # countls = pd.value_counts(df_month_loc["province"]).tolist()
    result = ""

    for org, count in zip(new_orgls[:3], new_countls[:3]):
        result = result + org + "（" + str(count) + "起）,"

    image4_text = (
        "图四解析："
        + minmonth
        + "至"
        + maxmonth
        + "，共"
        + str(len(new_orgls))
        + "家地区监管机构提出处罚意见，"
        + "排名前三的机构为："
        + result[: len(result) - 1]
    )
    st.markdown("#####  " + image4_text)

    # 图五解析：
    # get id list from searchdf
    # idlist = df_month["id"].tolist()
    # # get lawdetail
    # lawdf = get_lawcbirc()
    # # search lawdetail by selected_rows_id
    # selected_lawdetail = lawdf[lawdf["id"].isin(idlist)]

    # # law type count
    # lawtype = (
    #     selected_lawdetail.groupby("法律法规")["id"]
    #     .nunique()
    #     .reset_index(name="数量统计")
    # )
    # # sort by count
    # lawtype = lawtype.sort_values(by="数量统计", ascending=False)
    # x_data3 = lawtype["法律法规"].tolist()
    # y_data3 = lawtype["数量统计"].tolist()
    # bar3, lawtype_selected = print_bar(
    #     x_data3[:20], y_data3[:20], "处罚数量", "前20法律法规统计"
    # )

    # 图五解析开始
    # lawtype_count = lawtype[["法律法规", "数量统计"]]  # 把法律法规的数量进行统计
    # # pandas数据排序
    # lawtype_count = lawtype_count.sort_values("数量统计", ascending=False)
    # result5a = ""
    # for i in range(5):
    #     try:
    #         result5a = (
    #             result5a
    #             + str(lawtype_count.iloc[i, 0])
    #             + "("
    #             + str(lawtype_count.iloc[i, 1])
    #             + "起),"
    #         )
    #     except Exception as e:
    #         print(e)
    #         break
    # # st.markdown(
    # #     "##### 图五解析:法律法规统计-不同法规维度：处罚事件中，各违规类型中处罚数量排名前五分别为:"
    # #     + result5[: len(result5) - 1]
    # # )
    # # by具体条文
    # # lawdf["数量统计"] = ""
    # new_lawtype = (
    #     selected_lawdetail.groupby(["法律法规", "条文"])["id"]
    #     .nunique()
    #     .reset_index(name="数量统计")
    # )
    # # new_lawtype=lawdf.groupby(['法律法规','条文'])#%%%
    # new_lawtype["法律法规明细"] = (
    #     new_lawtype["法律法规"] + "(" + new_lawtype["条文"] + ")"
    # )

    # lawtype_count = new_lawtype[
    #     ["法律法规明细", "数量统计"]
    # ]  # 把法律法规的数量进行统计
    # # pandas数据排序
    # lawtype_count = lawtype_count.sort_values("数量统计", ascending=False)
    # result5b = ""
    # for i in range(5):
    #     try:
    #         result5b = (
    #             result5b
    #             + str(lawtype_count.iloc[i, 0])
    #             + "("
    #             + str(lawtype_count.iloc[i, 1])
    #             + "起),"
    #         )
    #     except Exception as e:
    #         print(e)
    #         break
    # image5_text = (
    #     " 图五解析:法律法规统计-不同法规维度：处罚事件中，各违规类型中处罚数量排名前五分别为:"
    #     + result5a[: len(result5a) - 1]
    #     + "\n"
    #     + "法律法规统计-具体条文维度：处罚事件中，各违规类型中处罚数量排名前五分别为:"
    #     + result5b[: len(result5b) - 1]
    # )
    # st.markdown("##### " + image5_text)

    # display summary
    # st.markdown("### 分析报告下载")

    # if st.button("生成分析报告"):
    #     t1 = time.localtime()
    #     t1 = time.strftime("%Y-%m-%d %H%M%S", t1)

    #     image1 = bar.render(path=os.path.join(pencbirc, t1 + "image1.html"))
    #     image2 = line.render(path=os.path.join(pencbirc, t1 + t1 + "image2.html"))
    #     image3 = map_data.render(path=os.path.join(pencbirc, t1 + t1 + "image3.html"))
    #     image4 = pie.render(path=os.path.join(pencbirc, t1 + t1 + "image4.html"))
    #     image5 = bar3.render(path=os.path.join(pencbirc, t1 + t1 + "image5.html"))
    #     # 做title
    #     title = st.session_state["keywords_cbirc"]
    #     title_str = ""
    #     title_str = "(分析范围：期间:" + str(title[0]) + "至" + str(title[1]) + ","
    #     if len(str(title[2])) != 0:
    #         title_str = title_str + "文号为:" + str(title[2]) + "，"
    #     if len(title[7]) != 0:
    #         title_str = title_str + "发文单位为:" + title[3] + "，"
    #     title_str = title_str[: len(title_str) - 1] + ")"
    #     title_str = "银保监处罚事件分析报告\n" + title_str

    #     file_name = make_docx(
    #         title_str,
    #         [image1_text, image2_text, image3_text, image4_text, image5_text],
    #         [image1, image2, image3, image4, image5],
    #     )
    #     st.download_button(
    #         "下载分析报告", data=file_name.read(), file_name="分析报告.docx"
    #     )


# display event detail
def display_eventdetail(search_df):
    # draw figure
    display_dfmonth(search_df)
    # get search result from session
    search_dfnew = st.session_state["search_result_cbirc"]
    total = len(search_dfnew)
    # st.sidebar.metric("总数:", total)
    st.markdown("### 搜索结果" + "(" + str(total) + "条)")
    # display download button
    st.download_button(
        "下载搜索结果",
        data=search_dfnew.to_csv().encode("utf_8_sig"),
        file_name="搜索结果.csv",
    )
    # display columns
    discols = ["id", "标题", "文号", "发布日期", "industry"]
    # get display df
    display_df = search_dfnew[discols]
    # change column name
    display_df.columns = ["id", "标题", "文号", "发布日期", "行业类型"]

    # reset index
    display_df.reset_index(drop=True, inplace=True)

    data = st.dataframe(display_df, on_select="rerun", selection_mode="single-row")

    selected_rows = data["selection"]["rows"]

    # data = df2aggrid(display_df)
    # display data
    # selected_rows = data["selected_rows"]
    if selected_rows == []:
        st.error("请先选择查看案例")
        st.stop()

    # get id column value from row number
    id = display_df.loc[selected_rows[0], "id"]

    # id = selected_rows[0]["id"]
    # select search_dfnew by id
    selected_rows_df = search_dfnew[search_dfnew["id"] == id]

    # select display columns
    selected_rows_df = selected_rows_df[
        [
            "发布日期",
            "summary",
            "wenhao",
            "people",
            "event",
            "law",
            "penalty",
            "org",
            "date",
            "category",
            "amount",
            "province",
            "industry",
            "内容",
        ]
    ]
    # rename columns
    selected_rows_df.columns = [
        "发布日期",
        "摘要",
        "文号",
        "当事人",
        "违法事实",
        "处罚依据",
        "处罚决定",
        "处罚机关",
        "处罚日期",
        "案件类型",
        "罚款金额",
        "处罚地区",
        "行业类型",
        "内容",
    ]

    # transpose and set column name
    selected_rows_df = selected_rows_df.astype(str).T
    # st.write(selected_rows_df)
    selected_rows_df.columns = ["内容"]
    # display selected rows by row with column name
    # for i in range(len(selected_rows_df)):
    #     st.markdown("##### " + selected_rows_df.index[i])
    #     st.markdown(selected_rows_df.iloc[i, 0])

    st.table(selected_rows_df)

    # get event detail url
    url = urldtl + str(id)
    # display url
    st.markdown("##### 案例链接")
    st.markdown(url)
    # get amtdf
    # amtdf = get_cbircamt()
    # # search amt by url
    # amtdata = amtdf[amtdf["id"] == id]
    # # display amount if amtdata is not empty
    # if amtdata.empty:
    #     st.error("没有找到相关罚款金额信息")
    # else:
    #     # display penalty amount
    #     amount = amtdata["amount"].values[0]
    #     st.metric("罚款金额", amount)

    # # get litigantdf
    # litigantdf = get_cbirclitigant()
    # # search litigant by id
    # litigantdata = litigantdf[litigantdf["id"] == id]
    # # display litigant if litigantdata is not empty
    # if litigantdata.empty:
    #     st.error("没有找到相关当事人信息")
    # else:
    #     # display litigant
    #     peoplels = litigantdata["peoplels"].values[0]
    #     orgls = litigantdata["orgls"].values[0]
    #     # convert list to string
    #     peoplestr = " ".join(peoplels)
    #     orgstr = " ".join(orgls)
    #     st.markdown("##### 当事人")
    #     st.markdown("个人 :" + peoplestr)
    #     st.markdown("机构 :" + orgstr)

    # # get labeldf
    # labeldf = get_cbirclabel()
    # # search labels by url
    # labeldata = labeldf[labeldf["id"] == id]
    # # display labels if labeldata is not empty
    # if labeldata.empty:
    #     st.error("没有找到相关标签")
    # else:
    #     # display labels
    #     labels = labeldata["labels"].values[0]
    #     scorels = labeldata["scores"].values[0]
    #     # convert scores to string
    #     scorels2 = ["%.3f" % x for x in scorels]
    #     scorestr = "/".join(scorels2)
    #     # st.markdown(scorestr)
    #     keywords = st_tags(
    #         label="##### 案件类型", text=scorestr, value=labels, suggestions=labels
    #     )

    # get lawdetail
    # lawdf = get_lawcbirc()
    # # search lawdetail by selected_rows_id
    # selected_rows_lawdetail = lawdf[lawdf["id"] == id]

    # if len(selected_rows_lawdetail) > 0:
    #     # display lawdetail
    #     st.markdown("##### 处罚依据")
    #     lawdata = selected_rows_lawdetail[["法律法规", "条文"]]
    #     # display lawdata
    #     for i in range(len(lawdata)):
    #         st.markdown(
    #             str(i + 1) + ": " + lawdata.iloc[i, 0] + " " + lawdata.iloc[i, 1]
    #         )

    # lawdtl = df2aggrid(lawdata)
    #     selected_law = lawdtl["selected_rows"]
    #     if selected_law == []:
    #         st.error("请先选择查看监管条文")
    #     else:
    #         # get selected_law's rule name
    #         selected_law_name = selected_law[0]["法律法规"]
    #         # get selected_law's rule article
    #         selected_law_article = selected_law[0]["条文"]
    #         # get law detail by name
    #         ruledf = get_rulelist_byname(selected_law_name, "", "", "", "")
    #         # get law ids
    #         ids = ruledf["lawid"].tolist()
    #         # get law detail by id
    #         metadf, dtldf = get_lawdtlbyid(ids)
    #         # display law meta
    #         st.write("监管法规")
    #         st.table(metadf)
    #         # get law detail by article
    #         articledf = dtldf[dtldf["标题"].str.contains(selected_law_article)]
    #         # display law detail
    #         st.write("监管条文")
    #         st.table(articledf)
    # else:
    #     st.write("没有相关监管法规")


# get sumeventdf in page number range
def get_sumeventdf(orgname, start, end):
    # choose orgname index
    org_index = {
        "银保监会机关": "4113",
        "银保监局本级": "4114",
        "银保监分局本级": "4115",
    }
    org_name_index = org_index[orgname]

    baseurl = "https://www.nfra.gov.cn/cbircweb/DocInfo/SelectDocByItemIdAndChild?itemId={}&pageSize=18&pageIndex={}"

    resultls = []
    errorls = []
    count = 0
    for i in range(start, end + 1):
        st.info("page: " + str(i))
        st.info(str(count) + " begin")
        url = baseurl.format(org_name_index, i)
        st.info("url:" + url)
        try:
            dd = requests.get(url, verify=False)
            sd = BeautifulSoup(dd.content, "html.parser")

            json_data = json.loads(str(sd.text))

            datals = json_data["data"]["rows"]

            df = pd.DataFrame(datals)
            resultls.append(df)
        except Exception as e:
            st.error("error!: " + str(e))
            errorls.append(url)

        mod = (i + 1) % 2
        if mod == 0 and count > 0:
            tempdf = pd.concat(resultls)
            savename = "tempsum-" + org_name_index + str(count + 1)
            savedf(tempdf, savename)

        wait = random.randint(2, 20)
        time.sleep(wait)
        st.info("finish: " + str(count))
        count += 1

    circdf = pd.concat(resultls)
    savecsv = "tempsumall" + org_name_index + str(count)
    savedf(circdf, savecsv)
    return circdf


def savedf(df, basename):
    savename = basename + ".csv"
    savepath = os.path.join(pencbirc, savename)
    df.to_csv(savepath)


# update sumeventdf
def update_sumeventdf(currentsum, orgname):
    org_name_index = org2name[orgname]
    # get sumeventdf
    oldsum = get_cbircsum(orgname)
    # get detail
    # oldsum = get_cbircdetail(orgname)
    if oldsum.empty:
        oldidls = []
    else:
        oldidls = oldsum["docId"].tolist()
    currentidls = currentsum["docId"].tolist()
    # print('oldidls:',oldidls)
    # print('currentidls:', currentidls)
    # get current idls not in oldidls
    newidls = [x for x in currentidls if x not in oldidls]
    # print('newidls:', newidls)
    # newidls=list(set(currentidls)-set(oldidls))
    newdf = currentsum[currentsum["docId"].isin(newidls)]
    # if newdf is not empty, save it
    if newdf.empty is False:
        newdf.reset_index(drop=True, inplace=True)
        nowstr = get_now()
        savename = "cbircsum" + org_name_index + nowstr
        savedf(newdf, savename)
        # save to update dtl list
        toupdname = "cbirctoupd" + org_name_index
        savedf(newdf, toupdname)
    return newdf


# update toupd
def update_toupd(orgname):
    org_name_index = org2name[orgname]
    # get sumeventdf
    currentsum = get_cbircsum(orgname)
    # get detail
    oldsum = get_cbircdetail(orgname)
    if oldsum.empty:
        oldidls = []
    else:
        oldidls = oldsum["id"].tolist()
    currentidls = currentsum["docId"].tolist()
    # get current idls not in oldidls
    newidls = [x for x in currentidls if x not in oldidls]
    newdf = currentsum[currentsum["docId"].isin(newidls)]
    # if newdf is not empty, save it
    if newdf.empty is False:
        newdf.reset_index(drop=True, inplace=True)
        # save to update dtl list
        toupdname = "cbirctoupd" + org_name_index
        savedf(newdf, toupdname)
    return newdf


# get current date and time string
def get_now():
    now = datetime.datetime.now()
    now_str = now.strftime("%Y%m%d%H%M%S")
    return now_str


# get current date and time string
def get_nowdate():
    now = datetime.datetime.now()
    now_str = now.strftime("%Y%m%d")
    return now_str


# get event detail
def get_eventdetail(eventsum, orgname):
    org_name_index = org2name[orgname]

    docidls = eventsum["docId"].tolist()

    # baseurl = (
    #     "https://www.cbirc.gov.cn/cn/static/data/DocInfo/SelectByDocId/data_docId="
    # )
    # update baseurl
    baseurl = "https://www.nfra.gov.cn/cn/static/data/DocInfo/SelectByDocId/data_docId"
    resultls = []
    errorls = []
    count = 0
    for i in docidls:
        st.info("id: " + str(i))
        st.info(str(count) + " begin")
        url = baseurl + "=" + str(i) + ".json"
        st.info("url:" + url)
        try:
            dd = requests.get(url, verify=False)
            
            # Check HTTP status code first
            if dd.status_code != 200:
                raise ValueError(f"HTTP {dd.status_code}: {dd.reason}")
            
            # Check content type
            content_type = dd.headers.get('content-type', '').lower()
            if 'application/json' not in content_type:
                st.warning(f"Unexpected content type: {content_type}")
            
            sd = BeautifulSoup(dd.content, "html.parser")
            response_text = str(sd.text).strip()
            
            # Check for common error responses
            if not response_text:
                raise ValueError("Empty response from server")
            
            if response_text.startswith('<') or 'nginx' in response_text.lower() or '404' in response_text:
                raise ValueError(f"Server returned HTML error page: {response_text[:100]}")
            
            # Try to parse JSON with better error handling
            try:
                json_data = json.loads(response_text)
            except json.JSONDecodeError as json_err:
                st.error(f"JSON decode error: {json_err}")
                st.error(f"Response text preview: {response_text[:200]}...")
                raise ValueError(f"Invalid JSON response: {json_err}")

            # Validate JSON structure
            if not isinstance(json_data, dict) or "data" not in json_data:
                raise ValueError("Invalid JSON structure: missing 'data' field")
            
            data_obj = json_data["data"]
            if not isinstance(data_obj, dict):
                raise ValueError("Invalid data structure in JSON response")
                
            title = data_obj.get("docTitle", "")
            subtitle = data_obj.get("docSubtitle", "")
            date = data_obj.get("publishDate", "")
            doc = data_obj.get("docClob", "")
            
            datals = {"title": title, "subtitle": subtitle, "date": date, "doc": doc}

            df = pd.DataFrame(datals, index=[0])
            df["id"] = i
            resultls.append(df)
        except Exception as e:
            st.error("error!: " + str(e))
            st.error("check url:" + url)
            errorls.append(i)

        mod = (count + 1) % 10
        if mod == 0 and count > 0:
            tempdf = pd.concat(resultls)
            savename = "tempdtl-" + org_name_index + str(count + 1)
            savedf(tempdf, savename)

        wait = random.randint(2, 20)
        time.sleep(wait)
        st.info("finish: " + str(count))
        count += 1
    # if resultls is not empty, save it
    if resultls:
        circdf = pd.concat(resultls)
        savecsv = "cbircdtl" + org_name_index + get_now()
        savedf(circdf, savecsv)
    else:
        circdf = pd.DataFrame()
    return circdf


# print bar graphs
# def print_bar(x_data, y_data, y_axis_name, title):
#     # draw echarts bar chart
#     bar = (
#         Bar()
#         .add_xaxis(xaxis_data=x_data)
#         .add_yaxis(series_name=y_axis_name, y_axis=y_data, yaxis_index=0)
#         .set_global_opts(
#             title_opts=opts.TitleOpts(title=title),
#             xaxis_opts=opts.AxisOpts(axislabel_opts=opts.LabelOpts(rotate=-15)),
#             visualmap_opts=opts.VisualMapOpts(max_=max(y_data), min_=min(y_data)),
#         )
#     )
#     # use events
#     events = {
#         "click": "function(params) { console.log(params.name); return params.name }",
#         # "dblclick":"function(params) { return [params.type, params.name, params.value] }"
#     }
#     # use events
#     clickevent = st_pyecharts(bar, events=events, height=400)
#     return bar, clickevent


def print_bar(x_data, y_data, y_axis_name, title):
    # Create a DataFrame from the input data
    data = pd.DataFrame({"月份": x_data, y_axis_name: y_data})
    # Create the bar chart
    fig = px.bar(
        data,
        x="月份",
        y=y_axis_name,
        title=title,
        color=y_axis_name,
        text=y_axis_name,
        color_continuous_scale=px.colors.sequential.Viridis,
    )

    # Display the chart
    event = st.plotly_chart(fig, use_container_width=True, on_select="rerun")

    monthselected = event["selection"]["point_indices"]

    if monthselected == []:
        clickevent = None
    else:
        clickevent = x_data[monthselected[0]]

    return fig, clickevent


# update the analysis data
# def update_cbircanalysis(orgname):
#     org_name_index = org2name[orgname]
#     # get cbirc detail
#     newdf = get_cbircdetail(orgname)
#     # get id list
#     newidls = newdf["id"].tolist()
#     # get cbirc analysis details
#     olddf = get_cbircanalysis(orgname)
#     # if olddf is not empty
#     if olddf.empty:
#         oldidls = []
#     else:
#         oldidls = olddf["id"].tolist()
#     # get new idls not in oldidls
#     updidls = [x for x in newidls if x not in oldidls]

#     upddf = newdf[newdf["id"].isin(updidls)]
#     # if newdf is not empty, save it
#     if upddf.empty is False:
#         splitdf = split_eventdoc(upddf)
#         updlen = len(splitdf)
#         st.info("拆分了" + str(updlen) + "条数据")
#         # combine with olddf
#         # upddf1 = pd.concat([splitdf, olddf])
#         upddf1 = splitdf
#         # # reset index
#         upddf1.reset_index(drop=True, inplace=True)
#         savename = "cbircanalysis" + org_name_index + get_now()
#         savedf(upddf1, savename)


# extract df by regex pattern matching
def extract_info(d1, pat):
    cols = ["标题", "文号", "发布日期", "doc1", "id"]
    ex1 = d1["doc1"].str.extract(pat)
    d2 = pd.concat([d1[cols], ex1], axis=1)
    r1 = d2[d2[0].notnull()]
    d3 = d2[d2[0].isnull()]
    return r1, d3


# split info from eventdf
# def split_eventdoc(d1):
#     # clean up
#     d1["doc1"] = d1["内容"].str.replace(r"\r|\n|\t|\xa0|\u3000|\s|\xa0", "")
#     # pat1 = r"行政处罚决定书文号(.*)被处罚当事人(.*)主要违法违规事实（案由）(.*)行政处罚依据(.*)行政处罚决定(.*)作出处罚决定的机关名称(.*)作出处罚决定的日期(.*)"
#     # r1, d2 = extract_info(d1, pat1)
#     # pat2 = r"行政处罚决定书文号(.*)被处罚当事人(.*)主要违法违规事实（案由）(.*)行政处罚依据(.*)行政处罚决定(.*)作出行政处罚决定的机关名称(.*)作出处罚决定的日期(.*)"
#     # r2, d3 = extract_info(d2, pat2)
#     # pat3 = r"行政处罚决定书文号(.*)被处罚当事人(.*)主要违法违规事实（案由）(.*)行政处罚依据(.*)行政处罚决定(.*)作出行政处罚决定的机关名称(.*)作出行政处罚决定的日期(.*)"
#     # r3, d4 = extract_info(d3, pat3)
#     # pat4 = r"行政处罚决定书文号(.*)被处罚当事人(.*)主要违法违规事实(.*)行政处罚依据(.*)行政处罚决定(.*)作出处罚决定的机关名称(.*)作出处罚决定的日期(.*)"
#     # r4, d5 = extract_info(d4, pat4)
#     # pat5 = r"处罚决定书文号(.*)被处罚当事人(.*)主要违法事实（案由）(.*)行政处罚依据(.*)行政处罚决定(.*)作出行政处罚的机关名称(.*)作出处罚决定的日期(.*)"
#     # r5, d6 = extract_info(d5, pat5)
#     # pat6 = (
#     #     r"行政处罚决定书文号(.*)被处罚当事人(.*)案由(.*)行政处罚依据(.*)行政处罚决定(.*)作出处罚决定的机关名称(.*)作出处罚决定的日期(.*)"
#     # )
#     # r6, d7 = extract_info(d6, pat6)

#     pat6 = r"(?:(?:行政处罚决定书文号|处罚决定书文号|行政处罚决定书文号|行政处罚决定文书号|行政处罚决定书号|行政处罚文书号|行政处罚决定文号)(.*?))(?:被处罚当事人|被处罚人|被处罚单位|单位名称|被处罚机构情况|被处罚个人)(.*?)(?:主要违法违规事实（案由）|主要违法违规事实|主要违法事实（案由）|案由|主要违法事实)(.*)行政处罚依据(.*)(?:行政处罚决定|行政处罚种类及金额|行政处理决定|行政处罚种类)(.*)(?:作出处罚决定的机关名称|作出行政处罚的机关名称|作出行政处罚决定的机关名称|做出处罚决定的机关名称|作出处罚的机关|作出处罚决定机关名称)(.*?)(?:(?:作出处罚决定的日期|作出行政处罚决定的日期|做出处罚决定的日期|出行政处罚决定的日期|作出处罚的日期)(.*))?$"
#     r6, d7 = extract_info(d1, pat6)

#     pat7 = r"([^，,：、]+罚[^，,：、]*?号.?)((?:被处罚|当事人|受处罚|姓名|名称|单位).*?)((?:现场检查|经检查|根据举报|。|经抽查|我|你|近期|一、|一、违法|一、经查|我局对|依据《|根据《|经查|我局自|我局于|我局在|依据有关法律|\d{4}年\d?\d月\d?\d日起|\d{4}年\d?\d月\d?\d日至|\d{4}年).*?。)((?:依据原《|上述|你公司|你作为|上述事实|综上|我局认为|上述未|上述行为|上述违法|根据|依据《|以上行为|该公司上述).*。)(.*?)((?:\d\d\d\d|.{4})年[^，,、]*日)"
#     r7, d8 = extract_info(d7, pat7)
#     pat8 = r"((?:被处罚|当事人：|受处罚|姓名).*?)((?:。|我|你|近期|一、违法|一、经查|我局对|依据《|根据《|经查|我局自|我局于|我局在|依据有关法律|\d{4}年\d?\d月\d?\d日起|\d{4}年\d?\d月\d?\d日至|\d{4}年\d?\d月\d?\d日).*?。)((?:你公司|你作为|上述事实|综上|我局认为|上述未|上述行为|上述违法|根据|依据《|以上行为|该公司上述).*。)(.*?)((?:\d\d\d\d|.{4})年[^，,、]*日)"
#     r8, d9 = extract_info(d8, pat8)
#     pat9 = r"(.*?)((?:我分局|我会|本会|经查|你公司|依据).*?。)([^。]*?(?:依据|我局|根据).*。)(中国[^。]*)(.{4}年.*)"
#     r9, d10 = extract_info(d9, pat9)

#     comdf = r6
#     [["标题", "文号", "发布日期", "doc1", "id", 0, 1, 2, 3, 4, 5, 6]]
#     comdfcols = [
#         "标题",
#         "文号",
#         "发布日期",
#         "doc1",
#         "id",
#         "行政处罚决定书文号",
#         "被处罚当事人",
#         "主要违法违规事实",
#         "行政处罚依据",
#         "行政处罚决定",
#         "作出处罚决定的机关名称",
#         "作出处罚决定的日期",
#     ]
#     comdf.columns = comdfcols
#     # st.write(comdf)

#     r7 = generate_lawls(r7)
#     # convert column list to string
#     r7["lawls"] = r7["lawls"].apply(lambda x: "，".join(x))
#     r7.columns = [
#         "标题",
#         "文号",
#         "发布日期",
#         "doc1",
#         "id",
#         "行政处罚决定书文号",
#         "被处罚当事人",
#         "主要违法违规事实",
#         "行政处罚决定",
#         "作出处罚决定的机关名称",
#         "作出处罚决定的日期",
#         "行政处罚依据",
#     ]
#     # st.write(r7)

#     r8 = generate_lawls(r8)
#     # convert column list to string
#     r8["lawls"] = r8["lawls"].apply(lambda x: "，".join(x))
#     r8.columns = [
#         "标题",
#         "文号",
#         "发布日期",
#         "doc1",
#         "id",
#         "被处罚当事人",
#         "主要违法违规事实",
#         "行政处罚决定",
#         "作出处罚决定的机关名称",
#         "作出处罚决定的日期",
#         "行政处罚依据",
#     ]
#     # st.write(r8)

#     r9 = generate_lawls(r9)
#     # convert column list to string
#     r9["lawls"] = r9["lawls"].apply(lambda x: "，".join(x))
#     r9.columns = [
#         "标题",
#         "文号",
#         "发布日期",
#         "doc1",
#         "id",
#         "被处罚当事人",
#         "主要违法违规事实",
#         "行政处罚决定",
#         "作出处罚决定的机关名称",
#         "作出处罚决定的日期",
#         "行政处罚依据",
#     ]
#     # st.write(r9)
#     st.markdown("### 拆分失败的数量: " + str(len(d10)))
#     r10 = d10[["标题", "文号", "发布日期", "doc1", "id"]]
#     r10 = generate_lawls(r10)
#     # convert column list to string if item is list

#     r10["lawls"] = r10["lawls"].apply(
#         lambda x: "，".join(x) if isinstance(x, list) else x
#     )
#     r10.columns = [
#         "标题",
#         "文号",
#         "发布日期",
#         "doc1",
#         "id",
#         "行政处罚依据",
#     ]
#     r10["主要违法违规事实"] = r10["doc1"]
#     # st.write(r10)
#     # combine all df
#     alldf = pd.concat([comdf, r7, r8, r9, r10])
#     return alldf


def generate_lawls(d1):
    compat = "(?!《).(《[^,，；。]*?》[^；。]*?第[^,，；。《]*条)"
    compat2 = "(?!《).(《[^,，；。]*?》)"

    d1["lawls"] = d1["doc1"].str.extractall(compat).groupby(level=0)[0].apply(list)
    d1["lawls"].fillna(
        d1["doc1"].str.extractall(compat2).groupby(level=0)[0].apply(list), inplace=True
    )
    return d1


# convert eventdf to lawdf
def generate_lawdf(d1):
    # clean up
    d1["doc1"] = d1["内容"].str.replace(r"\r|\n|\t|\xa0|\u3000|\s|\xa0", "")
    # compat = "(?!《).(《[^,，；。]*?》[^；。]*?第[^,，；。《]*条)"
    # compat2 = "(?!《).(《[^,，；。]*?》)"

    # d1["lawls"] = d1["doc1"].str.extractall(compat).groupby(level=0)[0].apply(list)
    # d1["lawls"].fillna(
    #     d1["doc1"].str.extractall(compat2).groupby(level=0)[0].apply(list), inplace=True
    # )
    d1 = generate_lawls(d1)
    d1["处理依据"] = d1["lawls"].fillna("").apply(lawls2dict)

    d2 = d1[["id", "处理依据"]]
    d3 = d2.explode("处理依据")
    d4 = d3["处理依据"].apply(pd.Series)
    d5 = pd.concat([d3, d4], axis=1)
    d6 = d5.explode("条文")
    # reset index
    d6.reset_index(drop=True, inplace=True)
    savedf(d6, "cbirclawdf")
    return d6


def lawls2dict(ls):
    try:
        result = []
        for item in ls:
            lawdict = dict()
            lawls = re.findall(r"《(.*?)》", item)
            #         print(lawls)
            artls = re.findall(r"(第[^《》、和章节款（）\(\)]*?条)", item)
            #         print(artls)
            lawdict["法律法规"] = lawls[0]
            lawdict["条文"] = artls
            result.append(lawdict)
        return result
    except Exception as e:
        st.error(str(e))
        return np.nan


def download_cbircsum(org_namels):
    st.markdown("#### 案例数据下载")

    for orgname in org_namels:
        st.markdown("##### " + orgname)
        # get orgname
        org_name_index = org2name[orgname]
        beginwith = "cbircsum" + org_name_index
        oldsum = get_csvdf(pencbirc, beginwith)
        lensum = len(oldsum)
        st.write("列表数据量: " + str(lensum))
        
        # Check if DataFrame is empty or doesn't have docId column
        if lensum > 0 and "docId" in oldsum.columns:
            # get unique id number
            idls = oldsum["docId"].unique()
            st.write("id数量: " + str(len(idls)))
        else:
            st.write("id数量: 0 (无数据或缺少docId列)")

        beginwith = "cbircdtl" + org_name_index
        dtl = get_csvdf(pencbirc, beginwith)
        lendtl = len(dtl)
        st.write("详情数据量: " + str(lendtl))
        
        # Check if DataFrame is empty or doesn't have id column
        if lendtl > 0 and "id" in dtl.columns:
            # get unique id number
            idls = dtl["id"].unique()
            st.write("id数量: " + str(len(idls)))
        else:
            st.write("id数量: 0 (无数据或缺少id列)")

        # Only process if data exists
        if lensum > 0 and "docId" in oldsum.columns:
            # drop duplicates
            oldsum.drop_duplicates("docId", inplace=True)
            # listname
            listname = "cbircsum" + org_name_index + get_nowdate() + ".csv"
            # download list data
            st.download_button(
                "下载列表数据", data=oldsum.to_csv().encode("utf_8_sig"), file_name=listname
            )
        else:
            st.warning("无列表数据可下载")

        # Only process if data exists
        if lendtl > 0 and "id" in dtl.columns:
            # drop duplicates
            dtl.drop_duplicates("id", inplace=True)
            # detailname
            detailname = "cbircdtl" + org_name_index + get_nowdate() + ".csv"
            # download detail data
            st.download_button(
                "下载详情数据", data=dtl.to_csv().encode("utf_8_sig"), file_name=detailname
            )
        else:
            st.warning("无详情数据可下载")

    st.markdown("#### 分类数据下载")

    # get analysis data
    # beginwith = "cbircanalysis" + org_name_index
    beginwith = "cbircsplit"
    analysis = get_csvdf(pencbirc, beginwith)
    lenanalysis = len(analysis)
    st.write("拆分数据量: " + str(lenanalysis))
    
    # Check if DataFrame is empty or doesn't have id column
    if lenanalysis > 0 and "id" in analysis.columns:
        # get unique id number
        idls = analysis["id"].unique()
        st.write("id数量: " + str(len(idls)))
        
        analysisname = "cbircsplit" + get_nowdate() + ".csv"
        # drop duplicates
        analysis.drop_duplicates("id", inplace=True)
        # download analysis data
        st.download_button(
            "下载拆分数据",
            data=analysis.to_csv().encode("utf_8_sig"),
            file_name=analysisname,
        )
    else:
        st.write("id数量: 0 (无数据或缺少id列)")
        st.warning("无拆分数据可下载")

    catdf = get_csvdf(pencbirc, "cbirccat")
    lencat = len(catdf)
    st.write("分类数据量: " + str(lencat))
    
    # Check if DataFrame is empty or doesn't have id column
    if lencat > 0 and "id" in catdf.columns:
        # get unique id number
        idls = catdf["id"].unique()
        st.write("id数量: " + str(len(idls)))
        
        allname = "cbirccat" + get_nowdate() + ".csv"
        # drop duplicates
        catdf.drop_duplicates("id", inplace=True)
        st.download_button(
            "下载分类数据", data=catdf.to_csv().encode("utf_8_sig"), file_name=allname
        )
    else:
        st.write("id数量: 0 (无数据或缺少id列)")
        st.warning("无分类数据可下载")


def get_cbirccat():
    amtdf = get_csvdf(pencbirc, "cbirccat")
    
    # process amount field - convert to float
    if "amount" in amtdf.columns:
        amtdf["amount"] = pd.to_numeric(amtdf["amount"], errors="coerce")
    
    # process province field - normalize province names
    if "province" in amtdf.columns:
        amtdf["province"] = amtdf["province"].apply(
            lambda x: normalize_province_name(str(x)) if pd.notna(x) and str(x).strip() else ""
        )
    
    # process industry field - clean and standardize
    if "industry" in amtdf.columns:
        amtdf["industry"] = amtdf["industry"].fillna("").astype(str)
    
    # process category field - clean and standardize
    if "category" in amtdf.columns:
        amtdf["category"] = amtdf["category"].fillna("").astype(str)
    
    # rename columns law to lawlist (if exists)
    if "law" in amtdf.columns:
        amtdf.rename(columns={"law": "lawlist"}, inplace=True)
    
    return amtdf


def sum_amount_by_month(df):
    # amtdf = get_cbircamt()
    # df1 = pd.merge(
    #     df, amtdf.drop_duplicates("id"), left_on="id", right_on="id", how="left"
    # )
    df1 = df
    df1["amount"] = df1["amount"].fillna(0)
    df1["发布日期"] = pd.to_datetime(df1["发布日期"]).dt.date
    # df=df[df['发文日期']>=pd.to_datetime('2020-01-01')]
    df1["month"] = df1["发布日期"].apply(lambda x: x.strftime("%Y-%m"))
    df_month_sum = df1.groupby(["month"])["amount"].sum().reset_index(name="sum")
    df_sigle_penalty = df1[["month", "amount"]]
    return df_month_sum, df_sigle_penalty

    # print line charts


# def print_line(x_data, y_data, y_axis_name, title):
#     # draw echarts line chart
#     line = (
#         Line()
#         .add_xaxis(x_data)
#         .add_yaxis(y_axis_name, y_data, label_opts=opts.LabelOpts(is_show=True))
#         .set_global_opts(
#             title_opts=opts.TitleOpts(title=title),
#             # legend_opts=opts.LegendOpts(pos_top="48%"),
#         )
#     )
#     # use events
#     events = {
#         "click": "function(params) { console.log(params.name); return params.name }",
#         # "dblclick":"function(params) { return [params.type, params.name, params.value] }"
#     }
#     # use events
#     clickevent = st_pyecharts(line, events=events, height=400)
#     return line, clickevent


def print_line(x_data, y_data, y_axis_name, title):
    # Create a DataFrame from the input data
    data = pd.DataFrame({"月份": x_data, y_axis_name: y_data})
    # Create the line chart
    fig = px.line(data, x="月份", y=y_axis_name, title=title, text=y_axis_name)

    # Display the chart
    event = st.plotly_chart(fig, use_container_width=True, on_select="rerun")

    monthselected = event["selection"]["point_indices"]

    if monthselected == []:
        clickevent = None
    else:
        clickevent = x_data[monthselected[0]]

    return fig, clickevent


def get_cbirclabel():
    labeldf = get_csvdf(pencbirc, "cbirclabel")
    # literal_eval apply to labels and scores
    # labeldf["labels"] = labeldf["labels"].apply(literal_eval)
    # labeldf["scores"] = labeldf["scores"].apply(literal_eval)
    # fillna
    labeldf = labeldf.fillna("")
    return labeldf[["id", "label"]]


def get_lawcbirc():
    lawdf = get_csvdf(pencbirc, "cbirclawdf")
    # fillna
    lawdf = lawdf.fillna("")
    return lawdf[["id", "处理依据", "法律法规", "条文"]]


def get_cbircloc():
    locdf = get_csvdf(pencbirc, "cbircloc")
    # fillna
    locdf = locdf.fillna("")
    return locdf[["id", "province", "city", "county"]]


# # province_name为省份名称列表；province_values为各省份对应值；title_name为标题,dataname为值标签（如：处罚案例数量）
# def print_map(province_name, province_values, title_name, dataname):
#     with open(mappath, "r", encoding="utf-8-sig") as f:
#         map = st_Map(
#             "china",
#             json.loads(f.read()),
#         )

#     map_data = (
#         Map()
#         .add(
#             dataname,
#             [list(z) for z in zip(province_name, province_values)],
#             "china",
#             is_roam=False,
#             is_map_symbol_show=False,
#         )
#         .set_series_opts(label_opts=opts.LabelOpts(is_show=True))
#         .set_global_opts(
#             title_opts=opts.TitleOpts(title=title_name),
#             visualmap_opts=opts.VisualMapOpts(
#                 max_=max(province_values)  # , range_color=["#F3F781", "#D04A02"]
#             ),
#         )
#     )
#     # st_pyecharts(map_data, map=map, height=700)  # ,width=800 )
#     return map_data, map


def print_map(province_name, province_values, title_name):
    # load the GeoJSON file
    china_geojson = json.load(open(mappath, "r", encoding="utf-8-sig"))

    # st.write(china_geojson)

    # Create a DataFrame from the provided data
    data = pd.DataFrame({"省份": province_name, "处罚数量": province_values})
    # Create the choropleth map
    # fig = px.choropleth(
    fig = px.choropleth_mapbox(
        data,
        geojson=china_geojson,
        featureidkey="properties.name",
        locations="省份",
        color="处罚数量",
        color_continuous_scale="Viridis",
        mapbox_style="carto-positron",
        zoom=2,
        center={"lat": 35, "lon": 105},
        # scope='asia',
        title=title_name,
    )

    # Add text labels
    fig.update_traces(
        text=data["处罚数量"],
    )

    # Update geos
    fig.update_geos(
        visible=False,
        fitbounds="locations",
    )

    # Update layout
    fig.update_layout(title_text=title_name, title_x=0.5)

    # Display the chart in Streamlit
    st.plotly_chart(fig, use_container_width=True)
    return fig


# combine count by province
def count_by_province(city_ls, count_ls):
    if len(city_ls) != len(count_ls):
        raise ValueError("城市列表和计数列表的长度必须相同")

    # Use Counter for efficient counting
    province_counts = Counter()

    for city, count in zip(city_ls, count_ls):
        province = normalize_province_name(city)
        if province and province != "未知省份":
            province_counts[province] += count
        else:
            province_counts["未知省份"] += count

    # Use sorted with key function for efficient sorting
    sorted_provinces = sorted(province_counts.items(), key=lambda x: (-x[1], x[0]))

    # Use zip for efficient unpacking
    if not sorted_provinces:
        return [], []
    provinces, counts = zip(*sorted_provinces)

    return list(provinces), list(counts)


# print pie charts
# def print_pie(namels, valuels, title):
#     pie = (
#         Pie()
#         .add(
#             "",
#             [list(z) for z in zip(namels, valuels)],
#             radius=["30%", "60%"],
#             # center=["35%", "50%"]
#         )
#         # set legend position
#         .set_global_opts(
#             title_opts=opts.TitleOpts(title=title)
#             # set legend position to down
#             ,
#             legend_opts=opts.LegendOpts(pos_bottom="bottom"),
#             visualmap_opts=opts.VisualMapOpts(max_=max(valuels)),
#         )
#         .set_series_opts(label_opts=opts.LabelOpts(formatter="{b}: {c}"))
#     )
#     events = {
#         "click": "function(params) { console.log(params.name); return params.name }",
#         # "dblclick":"function(params) { return [params.type, params.name, params.value] }"
#     }
#     clickevent = st_pyecharts(pie, events=events, height=650)  # width=800)
#     return pie, clickevent


def print_pie(namels, valuels, title):
    data = pd.DataFrame({"names": namels, "values": valuels})

    fig = px.pie(
        data,
        names="names",
        values="values",
        title=title,
        labels={"names": "名称", "values": "数量"},
    )
    fig.update_traces(textinfo="label+percent", insidetextorientation="radial")
    # Display the chart
    event = st.plotly_chart(fig, use_container_width=True, on_select="rerun")

    monthselected = event["selection"]["point_indices"]

    if monthselected == []:
        clickevent = None
    else:
        clickevent = namels[monthselected[0]]

    return fig, clickevent


def update_cbirclabel():
    # get cbirc detail
    newdf = get_cbircdetail("")
    # get id list
    newidls = newdf["id"].tolist()
    # display the num of new idls
    st.info("新案例数量：" + str(len(newidls)))
    # get amount details
    amtdf = get_cbirccat()
    # display the num of amtdf
    st.info("已分类案例数量：" + str(len(amtdf)))

    # get splitdf
    splitdf = get_cbircanalysis("")
    # display the num of splitdf
    st.info("已拆分案例数量：" + str(len(splitdf)))

    # if amtdf is not empty
    if amtdf.empty:
        amtoldidls = []
    else:
        amtoldidls = amtdf["id"].tolist()
    # get new idls not in oldidls
    amtupdidls = [x for x in newidls if x not in amtoldidls]

    amtupddf = newdf[newdf["id"].isin(amtupdidls)]
    # sory by date
    amtupddf = amtupddf.sort_values(by="发布日期", ascending=False)
    # reset index
    amtupddf.reset_index(drop=True, inplace=True)
    # display newdf
    st.markdown("### 待更新分类数据")
    st.write(amtupddf)
    # if newdf is not empty, save it
    if amtupddf.empty is False:
        updlen = len(amtupddf)
        st.info("待更新分类" + str(updlen) + "条数据")
        savename = "cbirc_tocat" + get_nowdate() + ".csv"
        # download detail data
        st.download_button(
            "下载分类案例数据",
            data=amtupddf.to_csv().encode("utf_8_sig"),
            file_name=savename,
        )
    else:
        st.info("无待更新分类数据")

    if splitdf.empty:
        splitoldidls = []
    else:
        splitoldidls = splitdf["id"].tolist()

    splitupdidls = [x for x in newidls if x not in splitoldidls]

    splitupddf = newdf[newdf["id"].isin(splitupdidls)]
    # reset index
    splitupddf.reset_index(drop=True, inplace=True)
    # display newdf
    st.markdown("### 待更新拆分数据")
    st.write(splitupddf)
    # if newdf is not empty, save it
    if splitupddf.empty is False:
        updlen = len(splitupddf)
        st.info("待更新拆分" + str(updlen) + "条数据")
        savename = "cbirc_tosplit" + get_nowdate() + ".csv"
        # download detail data
        st.download_button(
            "下载拆分案例数据",
            data=splitupddf.to_csv().encode("utf_8_sig"),
            file_name=savename,
        )
    else:
        st.info("无待更新拆分数据")

    # labelupddf = anadf[anadf["id"].isin(labelupdidls)]
    # # if newdf is not empty, save it
    # if labelupddf.empty is False:
    #     updlen = len(labelupddf)
    #     st.info("待更新标签" + str(updlen) + "条数据")
    #     savename = "cbirc_tolabel" + get_nowdate() + ".csv"
    #     # download detail data
    #     st.download_button(
    #         "下载案例数据",
    #         data=labelupddf.to_csv().encode("utf_8_sig"),
    #         file_name=savename,
    #     )

    # locupddf = anadf[anadf["id"].isin(locupdidls)]
    # # if newdf is not empty, save it
    # if locupddf.empty is False:
    #     updlen = len(locupddf)
    #     st.info("待更新地区" + str(updlen) + "条数据")
    #     savename = "cbirc_toloc" + get_nowdate() + ".csv"
    #     # download detail data
    #     st.download_button(
    #         "下载案例数据",
    #         data=locupddf.to_csv().encode("utf_8_sig"),
    #         file_name=savename,
    #     )

    # litigantupddf = anadf[anadf["id"].isin(litigantupdidls)]
    # # if newdf is not empty, save it
    # if litigantupddf.empty is False:
    #     updlen = len(litigantupddf)
    #     st.info("待更新当事人" + str(updlen) + "条数据")
    #     savename = "cbirc_tolitigant" + get_nowdate() + ".csv"
    #     # download detail data
    #     st.download_button(
    #         "下载案例数据",
    #         data=litigantupddf.to_csv().encode("utf_8_sig"),
    #         file_name=savename,
    #     )


def get_cbirclitigant():
    df = get_csvdf(pencbirc, "cbirclitigant")
    df["peoplels"] = df["peoplels"].apply(literal_eval)
    df["orgls"] = df["orgls"].apply(literal_eval)
    # fillna
    df = df.fillna("")
    cols = ["id", "peoplels", "orgls", "org"]
    df = df[cols]
    return df


def make_docx(
    title, text, image
):  # 制作docx的函数，title以str形式传入，其他以list的形式传入，输出为字符串的形式
    document = Document()

    # st.write(title_str)
    # add title
    document.add_paragraph().add_run(title).bold = True
    # document.add_paragraph(title)
    document.styles["Normal"].font.size = Pt(12)
    document.styles["Normal"].font.name = "Times New Roman"  # 设置西文字体
    document.styles["Normal"]._element.rPr.rFonts.set(
        qn("w:eastAsia"), "FangSong"
    )  # 设置中文字体使用字
    # document.styles['Normal'].font.bold = True
    # 加粗字体

    for i, j in zip(text, image):  # [image1_text,image2_text],[image1,image2]
        # document.styles['Normal'].font.bold = False
        t = time.localtime()
        t = time.strftime("%Y-%m-%d %H%M%S", t)
        make_snapshot(driver, j, t + ".png", is_remove_html=True)  #
        document.add_paragraph(i)
        document.styles["Normal"].font.size = Pt(12)
        document.styles["Normal"].font.name = "Times New Roman"  # 设置西文字体
        document.styles["Normal"]._element.rPr.rFonts.set(
            qn("w:eastAsia"), "FangSong"
        )  # 设置中文字体使用字体2->宋体
        document.add_picture(
            t + ".png", width=docx.shared.Inches(5.4)
        )  # 6英尺是最大宽度
        # print('当前图像高度', str(document.inline_shapes[0].height)+'当前图像宽度'+str(document.inline_shapes[0].width)) # 打印当前图片大小
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        os.remove(t + ".png")
    file_stream = io.BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    # document.save(t1+'.docx')
    return file_stream


def uplink_cbircsum():
    st.markdown("#### 案例数据上线")

    # get old sumeventdf
    eventdf = get_cbircdetail("")
    # get lengh
    oldlen = len(eventdf)
    st.write("案例数据量：" + str(oldlen))
    # get id nunique
    oldidn = eventdf["id"].nunique()
    st.write("案例数据id数：" + str(oldidn))
    # drop duplicate by id
    eventdf.drop_duplicates(subset=["id"], inplace=True)

    # download amount data (包含新增字段: amount, industry, category, province)
    amountdf = get_cbirccat()
    # get lengh
    amountlen = len(amountdf)
    st.write("分类数据量：" + str(amountlen))
    # get id nunique
    amountidn = amountdf["id"].nunique()
    st.write("分类数据id数：" + str(amountidn))
    # drop duplicate by id
    amountdf.drop_duplicates(subset=["id"], inplace=True)

    # download analysis data
    analysisdf = get_cbircanalysis("")
    # get lengh
    analysislen = len(analysisdf)
    st.write("分析数据量：" + str(analysislen))
    # get id nunique
    analysisidn = analysisdf["id"].nunique()
    st.write("分析数据id数：" + str(analysisidn))
    # drop duplicate by id
    analysisdf.drop_duplicates(subset=["id"], inplace=True)
    analysisname = "cbircanalysis" + get_nowdate() + ".csv"

    # get collection
    collection = get_collection("pencbirc", "cbircanalysis")

    # delete data from the MongoDB collection
    if st.button("删除案例数据"):
        delete_data(collection)
        st.success("案例数据删除成功！")

    # get all online data
    online_data = get_data(collection)

    # get unique id number from online data
    online_id = online_data["id"].nunique()

    # display online data
    st.write("在线案例数据量：" + str(online_id))

    # combine all data - 首先合并分析数据和事件数据
    alldf = pd.merge(analysisdf, eventdf, on="id", how="left")
    
    # 然后合并分类数据（包含新增字段）
    if not amountdf.empty:
        # 合并分类数据，包含amount, industry, category, province等新增字段
        alldf = pd.merge(alldf, amountdf, on="id", how="left")
        st.write("已合并分类数据，包含新增字段：amount, industry, category, province")
    
    # get different data
    diff_data = alldf[~alldf["id"].isin(online_data["id"])]

    # 准备上线数据的列，包含新增字段
    base_columns = [
        "标题",
        "文号", 
        "发布日期",
        "id",
        "wenhao",
        "people",
        "event",
        "law",
        "penalty",
        "org",
        "date",
    ]
    
    # 添加新增字段到列列表中（如果存在的话）
    additional_columns = []
    for col in ["amount", "industry", "category", "province"]:
        if col in diff_data.columns:
            additional_columns.append(col)
    
    # 选择要包含的列
    selected_columns = base_columns + additional_columns
    available_columns = [col for col in selected_columns if col in diff_data.columns]
    
    diff_data3 = diff_data[available_columns]
    
    # update column names - 基础列名映射
    column_mapping = {
        "标题": "标题",
        "文号": "文号",
        "发布日期": "发布日期",
        "id": "id",
        "wenhao": "行政处罚决定书文号",
        "people": "被处罚当事人",
        "event": "主要违法违规事实",
        "law": "行政处罚依据",
        "penalty": "行政处罚决定",
        "org": "作出处罚决定的机关名称",
        "date": "作出处罚决定的日期",
        # 新增字段保持原名，不进行重命名
        "amount": "amount",
        "industry": "industry", 
        "category": "category",
        "province": "province"
    }
    
    # 只重命名存在的列
    rename_dict = {k: v for k, v in column_mapping.items() if k in diff_data3.columns}
    diff_data3 = diff_data3.rename(columns=rename_dict)
    
    # filter out "" values
    # create a real copy to avoid chained assignment warnings
    diff_data4 = diff_data3[diff_data3["主要违法违规事实"].notnull()].copy()
    # convert date to datetime
    diff_data4.loc[:, "发布日期"] = pd.to_datetime(diff_data4["发布日期"])
    # fillna
    diff_data4 = diff_data4.fillna("")
    
    # 显示包含的新增字段信息
    if additional_columns:
        st.info(f"本次上线数据包含新增字段：{', '.join(additional_columns)}")

    # download different data
    st.download_button(
        "下载差异数据",
        data=diff_data4.to_csv().encode("utf_8_sig"),
        file_name=analysisname,
    )

    # display different data
    st.write("差异数据量：" + str(len(diff_data4)))
    st.write(diff_data4)

    # Insert data into the MongoDB collection
    if st.button("更新上线案例数据"):
        insert_data(diff_data4, collection)
        st.success("更新案例数据上线成功！包含新增字段数据已同步到MongoDB")
